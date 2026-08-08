import os
import re
from pathlib import Path

import fitz
import pdfplumber
from docx import Document
from langdetect import detect, LangDetectException


class DocumentParser:

    # ==========================================================
    # Supported Files
    # ==========================================================

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt"
    }

    # ==========================================================
    # Initialization
    # ==========================================================

    def __init__(self, file_path=None, text=None):

        self.file_path = (
            Path(file_path)
            if file_path
            else None
        )

        self.text_input = text

        self.extension = (
            self.file_path.suffix.lower()
            if self.file_path
            else None
        )

        self.doc = None

    # ==========================================================
    # File Validation
    # ==========================================================

    def validate_file(self):

        if self.file_path is None:
            raise ValueError("No file provided.")

        if not self.file_path.exists():
            raise FileNotFoundError(
                f"File not found: {self.file_path}"
            )

        if not self.file_path.is_file():
            raise ValueError(
                "Provided path is not a file."
            )

        if self.extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {self.extension}. "
                "Supported formats: PDF, DOCX and TXT."
            )

        if self.file_path.stat().st_size == 0:
            raise ValueError(
                "The uploaded file is empty."
            )

        return True

    # ==========================================================
    # Text Validation
    # ==========================================================

    def validate_text(self):

        if self.text_input is None:
            raise ValueError(
                "No text was provided."
            )

        if not self.text_input.strip():
            raise ValueError(
                "The pasted text is empty."
            )

        return True

    # ==========================================================
    # Text Cleaning
    # ==========================================================

    def clean_text(self, text):

        if not text:
            return ""

        # Normalize line endings
        text = text.replace(
            "\r\n",
            "\n"
        )

        text = text.replace(
            "\r",
            "\n"
        )

        # Remove excessive spaces/tabs
        text = re.sub(
            r"[ \t]+",
            " ",
            text
        )

        # Remove excessive blank lines
        text = re.sub(
            r"\n{2,}",
            "\n",
            text
        )

        # Remove spaces around line breaks
        text = re.sub(
            r" *\n *",
            "\n",
            text
        )

        return text.strip()

    # ==========================================================
    # PDF Extraction
    # ==========================================================

    def parse_pdf(self):

        try:

            self.doc = fitz.open(
                str(self.file_path)
            )

            if self.doc.page_count == 0:
                raise ValueError(
                    "The PDF contains no pages."
                )

            pages = []
            full_text = []

            for page_number, page in enumerate(
                self.doc,
                start=1
            ):

                raw_text = page.get_text("text")

                cleaned_text = self.clean_text(
                    raw_text
                )

                pages.append({
                    "page": page_number,
                    "text": cleaned_text,
                    "word_count": len(
                        cleaned_text.split()
                    ),
                    "character_count": len(
                        cleaned_text
                    )
                })

                full_text.append(
                    cleaned_text
                )

            return pages, "\n".join(full_text)

        except Exception as e:

            raise ValueError(
                f"Unable to read PDF: {e}"
            )

        finally:

            if self.doc is not None:
                self.doc.close()
                self.doc = None

    # ==========================================================
    # DOCX Extraction
    # ==========================================================

    def parse_docx(self):

        try:

            document = Document(
                str(self.file_path)
            )

            content = []

            # Paragraphs
            for paragraph in document.paragraphs:

                text = self.clean_text(
                    paragraph.text
                )

                if text:
                    content.append(text)

            # Tables
            for table in document.tables:

                for row in table.rows:

                    row_content = []

                    for cell in row.cells:

                        cell_text = self.clean_text(
                            cell.text
                        )

                        if cell_text:
                            row_content.append(
                                cell_text
                            )

                    if row_content:

                        content.append(
                            " | ".join(row_content)
                        )

            complete_text = self.clean_text(
                "\n".join(content)
            )

            pages = [{
                "page": 1,
                "text": complete_text,
                "word_count": len(
                    complete_text.split()
                ),
                "character_count": len(
                    complete_text
                )
            }]

            return pages, complete_text

        except Exception as e:

            raise ValueError(
                f"Unable to read DOCX: {e}"
            )

    # ==========================================================
    # TXT Extraction
    # ==========================================================

    def parse_txt(self):

        encodings = [
            "utf-8",
            "utf-8-sig",
            "cp1252",
            "latin-1"
        ]

        text = None

        for encoding in encodings:

            try:

                with open(
                    self.file_path,
                    "r",
                    encoding=encoding
                ) as file:

                    text = file.read()

                break

            except UnicodeDecodeError:
                continue

        if text is None:
            raise ValueError(
                "Unable to decode TXT file."
            )

        text = self.clean_text(text)

        pages = [{
            "page": 1,
            "text": text,
            "word_count": len(
                text.split()
            ),
            "character_count": len(text)
        }]

        return pages, text

    # ==========================================================
    # Pasted Text
    # ==========================================================

    def parse_text(self):

        self.validate_text()

        text = self.clean_text(
            self.text_input
        )

        pages = [{
            "page": 1,
            "text": text,
            "word_count": len(
                text.split()
            ),
            "character_count": len(text)
        }]

        return pages, text

    # ==========================================================
    # Metadata
    # ==========================================================

    def extract_metadata(self):

        if self.extension == ".pdf":

            try:

                doc = fitz.open(
                    str(self.file_path)
                )

                metadata = doc.metadata or {}

                result = {
                    "title": metadata.get(
                        "title"
                    ),
                    "author": metadata.get(
                        "author"
                    ),
                    "creator": metadata.get(
                        "creator"
                    ),
                    "producer": metadata.get(
                        "producer"
                    ),
                    "subject": metadata.get(
                        "subject"
                    ),
                    "keywords": metadata.get(
                        "keywords"
                    )
                }

                doc.close()

                return result

            except Exception:

                return {}

        elif self.extension == ".docx":

            try:

                document = Document(
                    str(self.file_path)
                )

                properties = (
                    document.core_properties
                )

                return {
                    "title": properties.title,
                    "author": properties.author,
                    "subject": properties.subject,
                    "keywords": properties.keywords,
                    "comments": properties.comments,
                    "last_modified_by":
                        properties.last_modified_by
                }

            except Exception:

                return {}

        # TXT and pasted text
        return {}

    # ==========================================================
    # Statistics
    # ==========================================================

    def get_statistics(
        self,
        text,
        pages
    ):

        words = len(
            text.split()
        )

        characters = len(text)

        paragraphs = len([
            p
            for p in text.split("\n")
            if p.strip()
        ])

        reading_time = max(
            1,
            round(words / 200)
        )

        return {
            "pages": len(pages),
            "words": words,
            "characters": characters,
            "paragraphs": paragraphs,
            "estimated_reading_minutes":
                reading_time
        }

    # ==========================================================
    # Language Detection
    # ==========================================================

    def detect_language(self, text):

        if not text or len(
            text.strip()
        ) < 20:

            return "unknown"

        try:

            return detect(text)

        except LangDetectException:

            return "unknown"

        except Exception:

            return "unknown"

    # ==========================================================
    # Document Type
    # ==========================================================

    def detect_document_type(self):

        if self.extension == ".pdf":

            try:

                doc = fitz.open(
                    str(self.file_path)
                )

                total_pages = (
                    doc.page_count
                )

                pages_with_text = 0

                for page in doc:

                    if page.get_text(
                        "text"
                    ).strip():

                        pages_with_text += 1

                doc.close()

                if pages_with_text == total_pages:
                    return "Searchable PDF"

                elif pages_with_text == 0:
                    return "Scanned / OCR Required"

                else:
                    return "Mixed PDF"

            except Exception:

                return "PDF"

        elif self.extension == ".docx":

            return "DOCX Document"

        elif self.extension == ".txt":

            return "Text Document"

        elif self.text_input is not None:

            return "Pasted Text"

        return "Unknown"

    # ==========================================================
    # File Information
    # ==========================================================

    def get_file_info(self):

        if self.file_path is None:

            return {
                "file_name": "Pasted Text",
                "file_extension": None,
                "file_size_bytes": None,
                "file_size_mb": None
            }

        size_bytes = (
            self.file_path.stat().st_size
        )

        return {
            "file_name":
                self.file_path.name,

            "file_extension":
                self.extension,

            "file_size_bytes":
                size_bytes,

            "file_size_mb":
                round(
                    size_bytes /
                    (1024 * 1024),
                    2
                )
        }

    # ==========================================================
    # Main Parse Function
    # ==========================================================

    def parse(
        self,
        extract_images=False,
        extract_tables=False
    ):

        # ------------------------------------------------------
        # Determine Input
        # ------------------------------------------------------

        if self.text_input is not None:

            pages, full_text = (
                self.parse_text()
            )

        else:

            self.validate_file()

            if self.extension == ".pdf":

                pages, full_text = (
                    self.parse_pdf()
                )

            elif self.extension == ".docx":

                pages, full_text = (
                    self.parse_docx()
                )

            elif self.extension == ".txt":

                pages, full_text = (
                    self.parse_txt()
                )

            else:

                raise ValueError(
                    "Unsupported document format."
                )

        # ------------------------------------------------------
        # Common Processing
        # ------------------------------------------------------

        statistics = self.get_statistics(
            full_text,
            pages
        )

        result = {

            "file_info":
                self.get_file_info(),

            "metadata":
                self.extract_metadata(),

            "statistics":
                statistics,

            "language":
                self.detect_language(
                    full_text
                ),

            "document_type":
                self.detect_document_type(),

            "pages":
                pages,

            "text":
                full_text,

            "images":
                [],

            "tables":
                []
        }

        # ------------------------------------------------------
        # Optional PDF Images
        # ------------------------------------------------------

        if (
            extract_images
            and self.extension == ".pdf"
        ):

            result["images"] = (
                self.extract_images()
            )

        # ------------------------------------------------------
        # Optional Tables
        # ------------------------------------------------------

        if extract_tables:

            result["tables"] = (
                self.extract_tables()
            )

        return result